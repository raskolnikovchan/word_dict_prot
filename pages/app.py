import docx.shared
import streamlit as st
import pandas as pd
import os
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth

# ユーザー設定読み込み
current_dir = os.getcwd()
yaml_path = os.path.join(current_dir, "config.yaml")

with open(yaml_path) as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    credentials=config['credentials'],
    cookie_name=config['cookie']['name'],
    cookie_key=config['cookie']['key'],
    cookie_expiry_days=config['cookie']['expiry_days'],
)

# 認証チェック
if not st.session_state.get("authentication_status"):
    st.warning('Please log in to access this page.')
    st.stop()

# ログイン成功
st.write("Profile Page Content")

# CSVファイルのパス
csv_path = "./pages/dict_words.csv"

def create_csv():
    if not os.path.exists(csv_path):
        # 初期データを持つ空のDataFrameを作成
        df = pd.DataFrame(columns=["name", "meaning"])
        df.to_csv(csv_path, index=False)

create_csv()

# セッションステートの初期化
if "word_list" not in st.session_state:
    st.session_state.word_list = []

if "new_words" not in st.session_state:
    st.session_state.new_words = []

if "doc_path" not in st.session_state:  
    st.session_state.doc_path = ""

if "concat_list" not in st.session_state:  
    st.session_state.concat_list = []

# 2. 単語を次々と登録していく
st.title("用語集作成")
st.write("## 1 単語を入力してください。全て入力したら完了ボタンを押してください。")
save_path = st.text_input("保存先のパスを入力してください", value=os.path.join(os.path.expanduser("~"), "Downloads"))
with st.form("word_form", clear_on_submit=True):
    name = st.text_input("用語名")
    add_button = st.form_submit_button("追加")

    if add_button and name:
        st.session_state.word_list.append(name)

st.write("現在の用語リスト:")
for i, word in enumerate(st.session_state.word_list):
    col1, col2 = st.columns([4,1])
    col1.write(f"{i}:{word}")
    if col2.button("削除", key=f"{i}_{word}"):
        st.session_state.word_list.remove(word)
        st.success(f"{word}を削除しました")



# 3. 登録された単語でCSVに存在しないものを抽出する
complete_button = st.button("完了")
if complete_button:
    df = pd.read_csv(csv_path)
    st.session_state.new_words = []

    for word in st.session_state.word_list:
        if (word not in df['name'].values) and (word not in st.session_state.new_words):
            st.session_state.new_words.append(word)

# 4. 存在しない単語の意味を手動で入力し、CSVに登録する
if st.session_state.new_words:
    st.write("### 1.5 以下の単語はデータベースに存在しません。意味を入力してください。")

    with st.form("meaning_form"):
        new_meaning = {}
        for i, word in enumerate(st.session_state.new_words):
            meaning = st.text_area(f"{i},{word} の意味を入力してください")
            new_meaning[word] = meaning
        
        submit_meaning = st.form_submit_button("登録")

        if submit_meaning:
            df = pd.read_csv(csv_path)

            for word, meaning in new_meaning.items():
                if meaning.strip():  # 空の意味が入力されていないか確認
                    # 新しい行を追加
                    new_row = pd.DataFrame({"name": [word], "meaning": [meaning]})
                    df = pd.concat([df, new_row], ignore_index=True)
            
            df.to_csv(csv_path, index=False)
            st.success("意味がデータベースに保存されました")

            # リセットする
            st.session_state.new_words = []

st.write("---")
st.write(f"## 2 全ての意味を入力したら以下のフォームからwordに出力してください。")

# Word出力のためのフォーム
with st.form("data_to_word", clear_on_submit=True):
    word_title = st.text_input("タイトル")
    not_index = st.checkbox("インデックスを表示しない")
    eliminate = st.checkbox("重複した単語を除外する")
    word_submit = st.form_submit_button("wordに出力する")
    
    if word_submit:
        doc = docx.Document()
        title_para = doc.add_paragraph(f"{word_title}")
        title_run = title_para.runs[0]
        title_run.font.size = docx.shared.Pt(30)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        df = pd.read_csv(csv_path)
        double_check = []
        for i, word in enumerate(st.session_state.word_list):
            if eliminate and (word in double_check):
                continue

            meaning_row = df[df['name'] == word]
            if not meaning_row.empty:
                meaning = meaning_row['meaning'].values[0]

                paragraph = doc.add_paragraph()
                if not_index:
                    pass
                else:
                    run_index = paragraph.add_run(f"{i} ")
                    run_index.font.size = docx.shared.Pt(16)

                # 太字にする
                run_word = paragraph.add_run(f"{word}：")
                run_word.font.bold = True  
                run_word.font.size = docx.shared.Pt(16)

                # 意味を追加
                run_meaning = paragraph.add_run(f"{meaning}")
                run_meaning.font.size = docx.shared.Pt(16)
                
                doc.add_paragraph()
                double_check.append(word)

        # ファイルを保存
        doc_path = f"{word_title}.docx"
        doc.save(doc_path)
        st.session_state.doc_path = doc_path  # セッションステートに保存
        st.success("データが保存されました。")
        st.session_state.word_list = []

# ダウンロードボタンを表示



st.write("--")
st.write("# 全てのデータをwordに保存する")
# Word出力のためのフォーム
with st.form("csv_to_word", clear_on_submit=True):
    word_title = st.text_input("タイトル")
    word_submit = st.form_submit_button("wordに出力する")
    if word_submit:
        doc = docx.Document()
        title_para = doc.add_paragraph(f"{word_title}")
        title_run = title_para.runs[0]
        title_run.font.size = docx.shared.Pt(30)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        df = pd.read_csv(csv_path)

        double_check = []
        for i, row in df.iterrows():
            word = row['name']
            meaning = row['meaning']

            if eliminate and (word in double_check):
                continue

            paragraph = doc.add_paragraph()
            run_index = paragraph.add_run(f"{i} ")
            run_index.font.size = docx.shared.Pt(16)

            # 太字にする
            run_word = paragraph.add_run(f"{word}：")
            run_word.font.bold = True  
            run_word.font.size = docx.shared.Pt(16)

            # 意味を追加
            run_meaning = paragraph.add_run(f"{meaning}")
            run_meaning.font.size = docx.shared.Pt(16)
            
            doc.add_paragraph()
            double_check.append(word)

        # ファイルを保存
        doc_path = f"{word_title}.docx"
        doc.save(doc_path)
        st.session_state.doc_path = doc_path  # セッションステートに保存
        st.success("データが保存されました。")
        st.session_state.word_list = []
        st.session_state.concat_list = []



#wordを結合して表示する
st.write("--")
st.write("## 結合word出力")
with st.form("concat_to_word", clear_on_submit=True):
    word_title = st.text_input("タイトル")
    eliminate = st.checkbox("重複した単語を除外する")
    word_submit = st.form_submit_button("wordに出力する")
    
    if word_submit:
        doc = docx.Document()
        title_para = doc.add_paragraph(f"{word_title}")
        title_run = title_para.runs[0]
        title_run.font.size = docx.shared.Pt(30)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        df = pd.read_csv(csv_path)
        double_check = []
        for word in st.session_state.concat_list:
            if eliminate and (word in double_check):
                continue

            meaning_row = df[df['name'] == word]
            if not meaning_row.empty:
                meaning = meaning_row['meaning'].values[0]

                paragraph = doc.add_paragraph()

                # 太字にする
                run_word = paragraph.add_run(f"{word}：")
                run_word.font.bold = True  
                run_word.font.size = docx.shared.Pt(16)

                # 意味を追加
                run_meaning = paragraph.add_run(f"{meaning}")
                run_meaning.font.size = docx.shared.Pt(16)
                
                doc.add_paragraph()
                double_check.append(word)

        # ファイルを保存
        doc_path = f"{word_title}.docx"
        doc.save(doc_path)
        st.session_state.doc_path = doc_path  # セッションステートに保存
        st.success("データが保存されました。")
        st.session_state.concat_list = []


if st.session_state.doc_path:  # ファイルが存在するか確認
    with open(st.session_state.doc_path, "rb") as f:
        st.download_button(
            label="ダウンロード",
            data=f,
            file_name=os.path.basename(st.session_state.doc_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


